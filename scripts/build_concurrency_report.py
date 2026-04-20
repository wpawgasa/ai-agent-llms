#!/usr/bin/env python3
"""Build docs/concurrency_latency_analysis.docx.

Reads all configs/models/cat_a/*.yaml, computes theoretical memory budgets and
predicted max-concurrent-sessions from known architecture parameters, then
writes a structured Word document.

Architecture parameters are hard-coded in ARCH_SPECS below.  Numbers come from
published technical reports and HuggingFace model cards; estimated values are
clearly annotated.  The doc includes a placeholder appendix for empirical
results produced by scripts/run_concurrency_benchmark.sh.

Usage:
    python3 scripts/build_concurrency_report.py
    # Output: docs/concurrency_latency_analysis.docx
"""

from __future__ import annotations

import math
from pathlib import Path

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

PROJECT_ROOT = Path(__file__).parent.parent
CONFIGS_DIR = PROJECT_ROOT / "configs" / "models" / "cat_a"
OUTPUT_PATH = PROJECT_ROOT / "docs" / "concurrency_latency_analysis.docx"

# H100 SXM 80GB hardware facts
H100_TOTAL_VRAM_GB = 80.0
H100_MEMORY_BW_TBPS = 3.35        # HBM3 memory bandwidth
H100_BF16_TFLOPS = 989.0
H100_FP8_TFLOPS = 1979.0
H100_SM_COUNT = 132

# vLLM defaults used in all Cat A configs
VLLM_GPU_MEM_UTIL = 0.90
VLLM_MAX_MODEL_LEN = 8192

# Usable VRAM (weight memory + KV cache must fit inside)
USABLE_VRAM_GB = H100_TOTAL_VRAM_GB * VLLM_GPU_MEM_UTIL  # 72 GB

# KV dtype memory multipliers (relative to BF16)
KV_DTYPE_MULTIPLIER = {
    "BF16": 1.0,
    "FP8": 0.5,
    "TurboQuant 3-bit (≈0.19×)": 3 / 16,   # 3 bits vs 16-bit BF16
    "TurboQuant 4-bit (≈0.25×)": 4 / 16,
}

# ---------------------------------------------------------------------------
# Architecture database
#
# Keys match model.name in the YAML configs.
# All KV cache formulas use:
#   kv_bytes_per_token = 2 × num_kv_layers × num_kv_heads × head_dim × bytes_per_elem
#
# For MLA (Multi-head Latent Attention): the compressed KV latent has dimension
# kv_lora_rank per layer, which replaces num_kv_heads × head_dim in the formula.
#
# For Mamba hybrid: only attention layers contribute to KV cache; Mamba SSM
# state is separate and fixed-size regardless of sequence length.
#
# Confidence: HIGH = from official HF config / paper; ESTIMATED = inferred from
# family architecture patterns.
# ---------------------------------------------------------------------------

ARCH_SPECS: dict[str, dict] = {
    "google/gemma-3-27b-it": {
        "display": "Gemma 3 27B-IT",
        "arch_type": "Dense GQA Transformer",
        "attention_detail": "GQA, alternating local (sliding-window 1024) + global layers",
        "num_kv_layers": 62,        # num_hidden_layers
        "num_kv_heads": 16,         # GQA 2:1 (32 Q / 16 KV)
        "head_dim": 256,            # Gemma 3 uses head_dim=256
        "moe": False,
        "mamba_hybrid": False,
        "mla": False,
        "turboquant_compatible": True,
        "turboquant_notes": "",
        "precision_notes": "BF16 weights only; no official FP8 checkpoint",
        "confidence": "HIGH",
        "source": "Gemma 3 Technical Report; HF config google/gemma-3-27b-it",
    },
    "google/gemma-4-26B-A4B-it": {
        "display": "Gemma 4 26B-A4B-IT (MoE, BF16)",
        "arch_type": "MoE GQA Transformer (256 experts, top-4 routing)",
        "attention_detail": "GQA with heterogeneous head_dim: global layers head_dim=256, local layers head_dim=128",
        "num_kv_layers": 46,
        "num_kv_heads": 8,          # ESTIMATED from Gemma4 family
        "head_dim": 256,            # global attention head_dim (dominant, forces TRITON_ATTN)
        "moe": True,
        "num_experts": 256,
        "top_k": 4,
        "params_active_b": 4,
        "mamba_hybrid": False,
        "mla": False,
        "turboquant_compatible": False,
        "turboquant_notes": "vLLM forces TRITON_ATTN (heterogeneous head_dim > 256 mixed with 128 local layers). "
                            "TRITON_ATTN does not support turboquant_* — see CLAUDE.md R8.",
        "precision_notes": "BF16 weights; use FP8 variant for lower VRAM",
        "confidence": "ESTIMATED",
        "source": "Gemma 4 architecture inferred from model card and vLLM source comments",
    },
    "RedHatAI/gemma-4-26B-A4B-it-FP8-Dynamic": {
        "display": "Gemma 4 26B-A4B-IT (MoE, FP8 weights)",
        "arch_type": "MoE GQA Transformer (256 experts, top-4 routing) — FP8 weights",
        "attention_detail": "Same as BF16 variant; FP8 is weight quantization (activations BF16)",
        "num_kv_layers": 46,
        "num_kv_heads": 8,
        "head_dim": 256,
        "moe": True,
        "num_experts": 256,
        "top_k": 4,
        "params_active_b": 4,
        "mamba_hybrid": False,
        "mla": False,
        "turboquant_compatible": False,
        "turboquant_notes": "Same TRITON_ATTN constraint as BF16 variant (R8); FP8 KV cache also blocked.",
        "precision_notes": "FP8 dynamic weight quantization halves weight VRAM (~26 GB); KV cache still BF16 or FP8",
        "confidence": "ESTIMATED",
        "source": "RedHatAI FP8-Dynamic checkpoint; architecture unchanged from BF16 base",
    },
    "google/gemma-4-31B-it": {
        "display": "Gemma 4 31B-IT (Dense, BF16)",
        "arch_type": "Dense GQA Transformer",
        "attention_detail": "GQA with heterogeneous head_dim (same family as 26B-A4B)",
        "num_kv_layers": 62,        # ESTIMATED: denser than 26B-A4B
        "num_kv_heads": 8,
        "head_dim": 256,
        "moe": False,
        "mamba_hybrid": False,
        "mla": False,
        "turboquant_compatible": False,
        "turboquant_notes": "vLLM forces TRITON_ATTN (heterogeneous head_dim) — same constraint as A4B variants (R8).",
        "precision_notes": "BF16 weights; ~62 GB VRAM for weights alone — very tight on H100 80GB",
        "confidence": "ESTIMATED",
        "source": "Inferred from YAML vram_estimate_gb=62 and Gemma4 architecture family",
    },
    "RedHatAI/gemma-4-31B-it-FP8-block": {
        "display": "Gemma 4 31B-IT (Dense, FP8 block weights)",
        "arch_type": "Dense GQA Transformer — FP8 block-quantized weights",
        "attention_detail": "Same architecture as BF16 variant; block FP8 for weights",
        "num_kv_layers": 62,
        "num_kv_heads": 8,
        "head_dim": 256,
        "moe": False,
        "mamba_hybrid": False,
        "mla": False,
        "turboquant_compatible": False,
        "turboquant_notes": "TRITON_ATTN constraint persists regardless of weight dtype (R8).",
        "precision_notes": "Block FP8 weights halve VRAM (~32 GB), freeing significant headroom for KV cache",
        "confidence": "ESTIMATED",
        "source": "RedHatAI FP8-block checkpoint; architecture unchanged from BF16 base",
    },
    "zai-org/GLM-4.7-Flash": {
        "display": "GLM-4.7-Flash (30B / 3.6B active MoE, MLA)",
        "arch_type": "MoE Transformer with Multi-head Latent Attention (MLA)",
        "attention_detail": "MLA: KV compressed to kv_lora_rank=512 per layer (vs full GQA KV space). "
                            "Similar to DeepSeek-V3 MLA design.",
        "num_kv_layers": 61,        # ESTIMATED from model size
        "num_kv_heads": None,       # MLA: KV is a latent vector, not per-head
        "head_dim": None,
        "kv_lora_rank": 512,        # MLA compressed KV dimension per layer (ESTIMATED)
        "moe": True,
        "mamba_hybrid": False,
        "mla": True,
        "turboquant_compatible": True,
        "turboquant_notes": "MLA compresses KV; TurboQuant would apply to the latent vector, "
                            "not standard K/V heads. Compatibility unverified.",
        "precision_notes": "BF16 weights; MLA means KV cache ~4× smaller than equivalent GQA model",
        "confidence": "ESTIMATED",
        "source": "GLM-4.7-Flash architectural design inferred from zai-org model card and MLA literature",
    },
    "mistralai/Mistral-Small-3.1-24B-Instruct-2503": {
        "display": "Mistral Small 3.1 24B-Instruct",
        "arch_type": "Dense GQA Transformer",
        "attention_detail": "GQA 4:1 (32 Q heads / 8 KV heads), sliding window attention",
        "num_kv_layers": 40,
        "num_kv_heads": 8,          # GQA 4:1
        "head_dim": 128,
        "moe": False,
        "mamba_hybrid": False,
        "mla": False,
        "turboquant_compatible": True,
        "turboquant_notes": "",
        "precision_notes": "BF16 weights; 48 GB — substantial KV cache headroom (24 GB+)",
        "confidence": "HIGH",
        "source": "Mistral Small 3.1 model card; HF config mistralai/Mistral-Small-3.1-24B-Instruct-2503",
    },
    "nvidia/Nemotron-3-Nano-30B-Instruct": {
        "display": "Nemotron-3-Nano 30B-Instruct (Mamba hybrid)",
        "arch_type": "Mamba-2 SSM + Sparse Attention Hybrid",
        "attention_detail": "~36 attention layers (with GQA 8 KV heads) alternating with Mamba-2 SSM layers. "
                            "Only attention layers contribute to KV cache.",
        "num_kv_layers": 36,        # ESTIMATED: ~60% Mamba, ~40% attention in hybrid schedule
        "num_kv_heads": 8,
        "head_dim": 128,
        "moe": True,                # 3.6B active out of 30B total
        "mamba_hybrid": True,
        "mla": False,
        "turboquant_compatible": False,
        "turboquant_notes": "TurboQuant boundary-layer protection requires uniform attention layers. "
                            "Mamba hybrid raises NotImplementedError at arg_utils.py:1650 — see CLAUDE.md R6.",
        "precision_notes": "vLLM compat uncertain for Mamba-2 hybrid (R6). HF generate() fallback may be needed.",
        "confidence": "ESTIMATED",
        "source": "Nemotron-3-Nano architecture from NVIDIA technical report; Mamba layer fraction estimated",
    },
    "Qwen/Qwen3-32B": {
        "display": "Qwen3-32B",
        "arch_type": "Dense GQA Transformer",
        "attention_detail": "GQA 8:1 (64 Q heads / 8 KV heads)",
        "num_kv_layers": 64,
        "num_kv_heads": 8,
        "head_dim": 128,
        "moe": False,
        "mamba_hybrid": False,
        "mla": False,
        "turboquant_compatible": True,
        "turboquant_notes": "",
        "precision_notes": "BF16 weights; 64 GB — only 8 GB KV cache headroom at GPU util 0.90. "
                           "FP8 KV dtype strongly recommended for meaningful concurrency.",
        "confidence": "HIGH",
        "source": "Qwen3 Technical Report; HF config Qwen/Qwen3-32B",
    },
    "Qwen/Qwen3.5-35B-A3B": {
        "display": "Qwen3.5-35B-A3B (DeltaNet + MoE, BF16)",
        "arch_type": "DeltaNet Linear Attention + MoE Hybrid",
        "attention_detail": "DeltaNet linear attention layers (hardware-efficient, O(1) state per layer) "
                            "interleaved with sparse standard attention layers. MoE FFN: ~3B active params.",
        "num_kv_layers": 16,        # ESTIMATED: only sparse attention layers have KV cache
        "num_kv_heads": 4,          # ESTIMATED from model family (DeltaNet uses fewer heads)
        "head_dim": 128,
        "moe": True,
        "mamba_hybrid": False,
        "mla": False,
        "turboquant_compatible": True,  # Only standard attention layers; DeltaNet has no KV cache
        "turboquant_notes": "TurboQuant applies only to standard attention layers. "
                            "DeltaNet state is separate. Verify boundary-layer handling.",
        "precision_notes": "BF16 weights OOM risk: 70 GB > 72 GB usable (R1). "
                           "Only viable with --max-model-len reduced or at gpu_memory_utilization < 0.90.",
        "confidence": "ESTIMATED",
        "source": "Qwen3.5 model card; DeltaNet architecture from Qwen team blog; KV layer count estimated",
    },
    "Qwen/Qwen3.6-35B-A3B": {
        "display": "Qwen3.6-35B-A3B (DeltaNet + MoE, BF16)",
        "arch_type": "DeltaNet Linear Attention + MoE Hybrid (next-gen Qwen3.5)",
        "attention_detail": "Same DeltaNet + MoE architecture as Qwen3.5 with expanded context (262K tokens). "
                            "Multi-Token Prediction (MTP) speculative decoding.",
        "num_kv_layers": 16,        # ESTIMATED: same family as Qwen3.5
        "num_kv_heads": 4,
        "head_dim": 128,
        "moe": True,
        "mamba_hybrid": False,
        "mla": False,
        "turboquant_compatible": True,
        "turboquant_notes": "Same as Qwen3.5 — only standard attention layers. "
                            "262K context makes max_model_len planning critical.",
        "precision_notes": "BF16 weights OOM risk: 70 GB > 72 GB usable (same as Qwen3.5, R1).",
        "confidence": "ESTIMATED",
        "source": "Qwen3.6 model card; extended-context variant of Qwen3.5-35B-A3B architecture",
    },
    "Qwen/Qwen3.6-35B-A3B-FP8": {
        "display": "Qwen3.6-35B-A3B (DeltaNet + MoE, FP8 weights)",
        "arch_type": "DeltaNet Linear Attention + MoE Hybrid — FP8 quantized weights",
        "attention_detail": "Same architecture as BF16 variant; FP8 weights halve VRAM to ~35 GB",
        "num_kv_layers": 16,
        "num_kv_heads": 4,
        "head_dim": 128,
        "moe": True,
        "mamba_hybrid": False,
        "mla": False,
        "turboquant_compatible": True,
        "turboquant_notes": "TurboQuant on standard attention layers. FP8 weights free ample KV cache headroom.",
        "precision_notes": "FP8 weights resolve the BF16 OOM risk; ~37 GB free for KV cache at GPU util 0.90",
        "confidence": "ESTIMATED",
        "source": "FP8 variant of Qwen3.6-35B-A3B; architecture unchanged from BF16",
    },
}

# ---------------------------------------------------------------------------
# Memory math helpers
# ---------------------------------------------------------------------------

BYTES_PER_ELEM = {"BF16": 2, "FP8": 1}


def kv_bytes_per_token(spec: dict, dtype: str = "BF16") -> float:
    """Return KV cache bytes per token for the given dtype.

    For MLA models: kv_bytes = 2 × num_kv_layers × kv_lora_rank × bytes_per_elem
    For standard GQA/GQA models: kv_bytes = 2 × num_kv_layers × num_kv_heads × head_dim × bytes_per_elem
    """
    bpe = BYTES_PER_ELEM.get(dtype, 2)
    if spec.get("mla"):
        lora_rank = spec.get("kv_lora_rank", 512)
        return 2 * spec["num_kv_layers"] * lora_rank * bpe
    return 2 * spec["num_kv_layers"] * spec["num_kv_heads"] * spec["head_dim"] * bpe


def max_concurrent_sessions(
    weight_vram_gb: float,
    kv_bytes_per_tok: float,
    context_length: int,
    available_vram_gb: float = USABLE_VRAM_GB,
) -> int:
    """Estimate maximum concurrent sessions that fit in memory.

    Uses the full context_length as the maximum KV cache size per session.
    In practice vLLM's paged allocator allows higher effective concurrency for
    shorter average contexts, so treat this as a conservative lower bound.
    """
    kv_budget_bytes = max(0.0, available_vram_gb - weight_vram_gb) * (1024**3)
    kv_per_session_bytes = kv_bytes_per_tok * context_length
    if kv_per_session_bytes <= 0:
        return 0
    return max(0, math.floor(kv_budget_bytes / kv_per_session_bytes))


# ---------------------------------------------------------------------------
# Docx helpers
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = tcPr.first_child_found_in("w:shd")
    if shd is None:
        from docx.oxml import OxmlElement
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)


def _bold_cell(cell) -> None:
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
        if para.runs:
            continue
        run = para.add_run(para.text)
        run.bold = True


def add_table(doc: Document, headers: list[str], rows: list[list[str]], style: str = "Table Grid") -> None:
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers), style=style)
    hdr_row = tbl.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        _set_cell_bg(cell, "2E75B6")
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx + 1]
        for c_idx, val in enumerate(row_data):
            row.cells[c_idx].text = str(val)
        bg = "DEEAF1" if r_idx % 2 == 0 else "FFFFFF"
        for cell in row.cells:
            _set_cell_bg(cell, bg)


# ---------------------------------------------------------------------------
# Report sections
# ---------------------------------------------------------------------------

def build_report(configs: list[dict]) -> Document:
    doc = Document()

    # Title
    title = doc.add_heading("Concurrency & Latency Analysis Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(
        "Category A LLM Models on NVIDIA H100 SXM 80GB with vLLM\n"
        "Analytical predictions — empirical results in Appendix C."
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    # -----------------------------------------------------------------------
    # Section 1: Executive Summary
    # -----------------------------------------------------------------------
    doc.add_heading("1. Executive Summary", 1)
    doc.add_paragraph(
        "This report analyses the maximum concurrent-request capacity for each of the 12 "
        "Category A models (15–35B parameters) served on a single NVIDIA H100 SXM 80GB GPU "
        "via vLLM. Concurrency is bounded primarily by KV cache VRAM after model weights are "
        "loaded. Key levers are:"
    )
    for bullet in [
        "KV cache dtype: BF16 (1×) → FP8 (2×) → TurboQuant 3-bit (~5.3×) compression.",
        "Weight dtype: FP8 checkpoint variants free 20–38 GB of additional VRAM.",
        "Architecture: MoE active params (3–4B) leave more headroom than dense models; "
        "MLA (GLM-4.7-Flash) compresses KV cache ~4× vs GQA baseline.",
        "Model-level vLLM constraints: Gemma4 (all variants) and Nemotron-3-Nano are "
        "TurboQuant-incompatible (see Section 5.2 and CLAUDE.md R6/R8).",
    ]:
        doc.add_paragraph(bullet, style="List Bullet")

    # Quick-reference summary table
    doc.add_paragraph()
    doc.add_paragraph("Predicted maximum concurrent sessions at ctx=4096 (conservative lower bound):")
    headers = ["Model", "Weight VRAM", "BF16 KV (ctx 4096)", "FP8 KV (ctx 4096)", "TQ 3-bit", "Notes"]
    summary_rows = _build_summary_rows(configs)
    add_table(doc, headers, summary_rows)

    # -----------------------------------------------------------------------
    # Section 2: H100 SXM 80GB Hardware Envelope
    # -----------------------------------------------------------------------
    doc.add_heading("2. H100 SXM 80GB Hardware Envelope", 1)
    doc.add_paragraph(
        "The NVIDIA H100 SXM 80GB is the target hardware for all experiments. "
        "Key specifications that constrain LLM serving throughput and concurrency:"
    )
    hw_rows = [
        ["Total VRAM", "80 GB HBM3"],
        ["Memory bandwidth", "3.35 TB/s"],
        ["BF16 Tensor Core throughput", "989 TFLOPS"],
        ["FP8 Tensor Core throughput", "1,979 TFLOPS (2× BF16)"],
        ["NVLink bandwidth", "900 GB/s bidirectional (multi-GPU, N/A single-GPU)"],
        ["SM count", "132 Streaming Multiprocessors"],
        ["Usable VRAM at gpu_memory_utilization=0.90", f"{USABLE_VRAM_GB:.0f} GB"],
    ]
    add_table(doc, ["Parameter", "Value"], hw_rows)

    doc.add_paragraph()
    doc.add_paragraph(
        "Decode-phase bottleneck: For models >15B parameters, the bottleneck during the "
        "auto-regressive decode phase is memory bandwidth, not compute.  Each decode step "
        "reads the entire KV cache plus model weights from HBM — arithmetic intensity falls "
        "below the BF16 compute/bandwidth ridge point (~74 FLOPS/byte) for batch sizes "
        "below ~500 concurrent tokens.  Increasing concurrency raises arithmetic intensity "
        "and improves hardware utilisation until the memory capacity (not bandwidth) becomes "
        "the ceiling."
    )
    doc.add_paragraph(
        "Prefill-phase: For long prompts (>1K tokens) the workload is compute-bound on H100.  "
        "TTFT is dominated by prompt processing latency, which scales linearly with prompt "
        "length and inversely with batch throughput."
    )

    # -----------------------------------------------------------------------
    # Section 3: vLLM Scheduling Primer
    # -----------------------------------------------------------------------
    doc.add_heading("3. vLLM Scheduling and KV Cache Quantisation", 1)
    doc.add_heading("3.1  PagedAttention and Continuous Batching", 2)
    doc.add_paragraph(
        "vLLM allocates KV cache in fixed-size pages (default 16 tokens/block).  "
        "Sequences share the physical KV block pool — no over-allocation, no fragmentation.  "
        "Continuous batching inserts new requests into decode batches as soon as a slot "
        "frees, maximising GPU utilisation.  Maximum in-flight sequences is limited by "
        "the KV block pool size, which depends on:"
    )
    for item in [
        "gpu_memory_utilization × VRAM_total − model_weight_memory = KV block pool",
        "block_size (tokens per block) × bytes_per_block = KV block size",
        "max_model_len controls the maximum KV cache a single sequence may consume",
    ]:
        doc.add_paragraph(item, style="List Bullet")
    doc.add_paragraph(
        "All Cat A configs set gpu_memory_utilization=0.90 (72 GB usable) and "
        "max_model_len=8192.  The benchmark script overrides max_model_len at launch to "
        "accommodate the largest requested context length plus output tokens."
    )

    doc.add_heading("3.2  KV Cache Dtype Options and Memory Multipliers", 2)
    dtype_rows = [
        ["auto / BF16", "1.0×", "2 bytes per element", "Full precision KV; most headroom consumed"],
        ["FP8 (--kv-cache-dtype fp8)", "0.50×", "1 byte per element", "Native vLLM support; ~0–1% quality drop; 2× sessions"],
        ["KIVI 2-bit", "0.125×", "0.25 bytes per element", "Asymmetric per-channel; no calibration needed"],
        ["KIVI 4-bit", "0.25×", "0.5 bytes per element", "Higher quality than 2-bit; same no-calibration property"],
        ["KVQuant NUQ", "0.125–0.25×", "variable", "NUQ codebooks; calibration required; dense-sparse decomp"],
        ["TurboQuant 3-bit", "≈0.19×", "3 bits per elem (packed)", "Codebook + rotation + QJL residual; ~5.3× vs BF16"],
        ["TurboQuant 4-bit", "≈0.25×", "4 bits per elem", "Higher quality than 3-bit variant"],
    ]
    add_table(
        doc,
        ["Dtype", "Memory multiplier", "Storage", "Notes"],
        dtype_rows,
    )
    doc.add_paragraph()
    doc.add_paragraph(
        "TurboQuant incompatibility: Gemma4 (all variants) forces TRITON_ATTN in vLLM due to "
        "heterogeneous head dimensions (head_dim mixes 256 and 128 across global/local layers).  "
        "TRITON_ATTN does not accept turboquant_* dtypes.  Nemotron-3-Nano's Mamba hybrid "
        "raises NotImplementedError in TurboQuant's boundary-layer protection.  "
        "Both model families are limited to FP8 / KIVI / KVQuant for KV compression."
    )

    doc.add_heading("3.3  Benchmark Degradation Policy", 2)
    doc.add_paragraph(
        "The empirical benchmark (scripts/run_concurrency_benchmark.sh) defines "
        "'no degradation' as the highest concurrency level N satisfying both:"
    )
    doc.add_paragraph(
        "  (1)  TTFT p95[N]  ≤  2.0 × TTFT p95[concurrency=1]",
        style="List Number",
    )
    doc.add_paragraph(
        "  (2)  Request failure rate[N]  ≤  1 %",
        style="List Number",
    )
    doc.add_paragraph(
        "The full concurrency curve is recorded (not a binary search), so latency vs. "
        "concurrency plots can be generated.  Context lengths swept: 2048, 4096, 8192 tokens."
    )

    # -----------------------------------------------------------------------
    # Section 4: Per-Model Analysis
    # -----------------------------------------------------------------------
    doc.add_heading("4. Per-Model Analysis", 1)
    doc.add_paragraph(
        "Each sub-section presents the architecture card, KV cache memory formula, "
        "available VRAM budget, and predicted maximum concurrent sessions at ctx=4096 "
        "for BF16 and FP8 KV dtypes.  Numbers are analytical lower bounds; vLLM's "
        "paged allocator allows higher effective concurrency when average session "
        "context is shorter than the maximum."
    )

    for cfg in configs:
        _write_model_section(doc, cfg)

    # -----------------------------------------------------------------------
    # Section 5: Cross-Model Comparison Tables
    # -----------------------------------------------------------------------
    doc.add_heading("5. Cross-Model Comparison", 1)
    doc.add_heading("5.1  KV Cache Budget and Predicted Concurrency Matrix", 2)
    doc.add_paragraph(
        "Table below: weight VRAM from YAML vram_estimate_gb; KV budget = "
        f"{USABLE_VRAM_GB:.0f} GB − weight_VRAM; max concurrent sessions at ctx=4096 "
        "= floor(kv_budget / kv_per_session).  Values are conservative lower bounds."
    )
    matrix_headers = [
        "Model", "Weight VRAM (GB)", "KV budget (GB)",
        "BF16 KV/token (KB)", "Max sessions BF16", "Max sessions FP8 KV",
        "Max sessions TQ-3bit",
    ]
    matrix_rows = _build_matrix_rows(configs)
    add_table(doc, matrix_headers, matrix_rows)

    doc.add_heading("5.2  TurboQuant Compatibility Matrix", 2)
    tq_headers = ["Model", "TQ Compatible?", "Reason"]
    tq_rows = []
    for cfg in configs:
        spec = ARCH_SPECS.get(cfg["model"]["name"], {})
        compat = "✓ Yes" if spec.get("turboquant_compatible") else "✗ No"
        reason = spec.get("turboquant_notes", "—") or "Compatible"
        tq_rows.append([spec.get("display", cfg["model"]["name"]), compat, reason])
    add_table(doc, tq_headers, tq_rows)

    # -----------------------------------------------------------------------
    # Section 6: Benchmark Methodology
    # -----------------------------------------------------------------------
    doc.add_heading("6. Empirical Benchmark Methodology", 1)
    doc.add_paragraph(
        "Section 4's predictions are based on analytical memory formulas and published "
        "architecture parameters.  The empirical benchmark measures actual TTFT, TPOT, "
        "ITL, throughput, and peak VRAM under real load, accounting for vLLM's scheduling "
        "overhead, CUDA kernel launch latency, and prefill/decode pipeline effects."
    )
    doc.add_heading("6.1  Script", 2)
    doc.add_paragraph("scripts/run_concurrency_benchmark.sh — single model, single dtype.")
    para = doc.add_paragraph()
    run = para.add_run(
        "./scripts/run_concurrency_benchmark.sh configs/models/cat_a/qwen3_32b.yaml \\\n"
        "  --kv-cache-dtype fp8 \\\n"
        "  --context-lengths 2048,4096,8192 \\\n"
        "  --concurrency-levels 1,2,4,8,16,32,64,128,256,512,1024 \\\n"
        "  --requests-per-level 64 \\\n"
        "  --output-tokens 128"
    )
    run.font.name = "Courier New"
    run.font.size = Pt(9)

    doc.add_heading("6.2  Metrics Captured Per Level", 2)
    metric_rows = [
        ["TTFT p50/p95/p99 (ms)", "Time from request submit to first output token"],
        ["TPOT p50/p95/p99 (ms)", "(E2E − TTFT) / (output_tokens − 1): per-token decode time"],
        ["ITL p50/p95/p99 (ms)", "Inter-token latency for each successive output token"],
        ["E2E latency p50/p95/p99 (ms)", "Total request round-trip time"],
        ["Throughput (output tok/s)", "Aggregate output tokens / wall-clock time for the level"],
        ["Goodput (req/s)", "Successful requests / wall-clock time"],
        ["Success rate", "Fraction of requests completing without error"],
        ["Peak VRAM (GB)", "Maximum GPU memory observed during the level (nvidia-smi polling)"],
    ]
    add_table(doc, ["Metric", "Definition"], metric_rows)

    doc.add_heading("6.3  Prompt Shape", 2)
    doc.add_paragraph(
        "Input: ~512 tokens (business-workflow description, looped to fill the target size).  "
        "Output: max_tokens=128.  Temperature=0.0 (greedy; maximises reproducibility).  "
        "Warmup: 8 requests discarded per level to allow KV cache warm-up."
    )

    # -----------------------------------------------------------------------
    # Section 7: Risks and Known Deviations
    # -----------------------------------------------------------------------
    doc.add_heading("7. Risks and Known Deviations", 1)
    risk_rows = [
        ["R1", "Qwen3.5/3.6-35B-A3B BF16 OOM",
         "70 GB weights > 72 GB usable; only viable at reduced max_model_len or lower gpu_memory_utilization",
         "Use FP8 variant (Qwen3.6-35B-A3B-FP8) or reduce max_model_len to 2048"],
        ["R6", "Nemotron-3-Nano vLLM incompatibility",
         "Mamba-2 hybrid may not be supported by vLLM's attention backend for all dtypes",
         "HF generate() fallback; skip TurboQuant dtypes; benchmark with --kv-cache-dtype auto only"],
        ["R8", "Gemma4 + TurboQuant incompatible",
         "vLLM forces TRITON_ATTN for heterogeneous head_dim; TRITON_ATTN rejects turboquant_* dtype",
         "Use FP8 / KIVI / KVQuant for all Gemma4 cells in the benchmark matrix"],
        ["R-MLA", "GLM-4.7-Flash MLA KV format",
         "MLA's compressed KV latent may not be compatible with all KV dtype options in vLLM",
         "Benchmark with auto and fp8 first; verify MLA passthrough for TurboQuant"],
        ["R-OOM2", "Qwen3-32B BF16 KV headroom",
         "Only 8 GB KV budget at BF16; max_model_len must be small to allow >1 concurrent session",
         "Always benchmark Qwen3-32B with --kv-cache-dtype fp8 or turboquant_3bit_nc"],
    ]
    add_table(doc, ["ID", "Risk", "Description", "Mitigation"], risk_rows)

    # -----------------------------------------------------------------------
    # Appendix C: Empirical Results Placeholder
    # -----------------------------------------------------------------------
    doc.add_page_break()
    doc.add_heading("Appendix C: Empirical Results (TBD)", 1)
    doc.add_paragraph(
        "This appendix is a placeholder.  Populate rows by running the benchmark script "
        "for each (model, dtype) pair and aggregating results/concurrency/*.json."
    )
    doc.add_paragraph(
        "Command to run all Cat A models at fp8 dtype (example):",
    )
    para = doc.add_paragraph()
    run = para.add_run(
        "for cfg in configs/models/cat_a/*.yaml; do\n"
        "  ./scripts/run_concurrency_benchmark.sh \"$cfg\" --kv-cache-dtype fp8\n"
        "done"
    )
    run.font.name = "Courier New"
    run.font.size = Pt(9)

    doc.add_paragraph()
    emp_headers = [
        "Model", "KV dtype", "ctx=2048 max conc.", "ctx=4096 max conc.", "ctx=8192 max conc.",
        "TTFT p95 @baseline (ms)", "Peak VRAM (GB)",
    ]
    emp_rows = [
        [spec.get("display", name), "fp8", "TBD", "TBD", "TBD", "TBD", "TBD"]
        for name, spec in ARCH_SPECS.items()
    ]
    add_table(doc, emp_headers, emp_rows)

    return doc


def _write_model_section(doc: Document, cfg: dict) -> None:
    model_name = cfg["model"]["name"]
    spec = ARCH_SPECS.get(model_name, {})
    display = spec.get("display", model_name)
    weight_gb = cfg["model"].get("vram_estimate_gb", 0)

    doc.add_heading(f"4.x  {display}", 2)

    # Architecture card
    arch_rows = [
        ["Full model name", model_name],
        ["Architecture type", spec.get("arch_type", "—")],
        ["Attention detail", spec.get("attention_detail", "—")],
        ["KV-contributing layers", str(spec.get("num_kv_layers", "—"))],
    ]
    if not spec.get("mla"):
        arch_rows += [
            ["KV heads per layer", str(spec.get("num_kv_heads", "—"))],
            ["Head dimension", str(spec.get("head_dim", "—"))],
        ]
    else:
        arch_rows.append(["MLA kv_lora_rank", str(spec.get("kv_lora_rank", "—"))])
    arch_rows += [
        ["MoE", "Yes" if spec.get("moe") else "No"],
        ["Mamba hybrid", "Yes" if spec.get("mamba_hybrid") else "No"],
        ["MLA", "Yes" if spec.get("mla") else "No"],
        ["TurboQuant compatible", "Yes" if spec.get("turboquant_compatible") else "No"],
        ["Weight VRAM (BF16)", f"{weight_gb} GB"],
        ["Confidence", spec.get("confidence", "—")],
        ["Source", spec.get("source", "—")],
    ]
    add_table(doc, ["Parameter", "Value"], arch_rows)

    # Memory math
    doc.add_paragraph()
    kv_budget = max(0.0, USABLE_VRAM_GB - weight_gb)
    doc.add_paragraph(
        f"KV cache budget at gpu_memory_utilization=0.90:  "
        f"{USABLE_VRAM_GB:.0f} GB − {weight_gb} GB = {kv_budget:.1f} GB"
    )

    if spec and spec.get("num_kv_layers"):
        for dtype_label, mult in [("BF16", 1.0), ("FP8", 0.5)]:
            try:
                kv_bpt = kv_bytes_per_token(spec, "BF16") * mult
                kv_bpt_kb = kv_bpt / 1024
                for ctx in (2048, 4096, 8192):
                    sessions = max_concurrent_sessions(weight_gb, kv_bpt, ctx)
                    doc.add_paragraph(
                        f"  KV dtype {dtype_label}: {kv_bpt_kb:.1f} KB/token × {ctx} tokens = "
                        f"{kv_bpt * ctx / 1024**3 * 1000:.0f} MB/session → "
                        f"max {sessions} concurrent sessions"
                    )
            except Exception:
                doc.add_paragraph(f"  KV dtype {dtype_label}: insufficient architecture data for formula")

    # Notes
    for note_key in ("turboquant_notes", "precision_notes"):
        note = spec.get(note_key, "")
        if note:
            p = doc.add_paragraph()
            run = p.add_run(f"Note: {note}")
            run.italic = True

    doc.add_paragraph()


def _build_summary_rows(configs: list[dict]) -> list[list[str]]:
    rows = []
    for cfg in configs:
        name = cfg["model"]["name"]
        spec = ARCH_SPECS.get(name, {})
        weight_gb = cfg["model"].get("vram_estimate_gb", 0)
        display = spec.get("display", name)
        compat = "✓" if spec.get("turboquant_compatible") else "✗ R8/R6"
        if spec and spec.get("num_kv_layers"):
            try:
                kv_bf16 = kv_bytes_per_token(spec, "BF16")
                kv_fp8 = kv_bytes_per_token(spec, "FP8")
                s_bf16 = max_concurrent_sessions(weight_gb, kv_bf16, 4096)
                s_fp8 = max_concurrent_sessions(weight_gb, kv_fp8, 4096)
                tq_mult = 3 / 16
                s_tq = max_concurrent_sessions(weight_gb, kv_bf16 * tq_mult, 4096) if spec.get("turboquant_compatible") else "N/A"
                rows.append([display, f"{weight_gb} GB", str(s_bf16), str(s_fp8), str(s_tq), ""])
            except Exception:
                rows.append([display, f"{weight_gb} GB", "—", "—", "—", "Arch data incomplete"])
        else:
            rows.append([display, f"{weight_gb} GB", "—", "—", "—", "Arch data incomplete"])
    return rows


def _build_matrix_rows(configs: list[dict]) -> list[list[str]]:
    rows = []
    for cfg in configs:
        name = cfg["model"]["name"]
        spec = ARCH_SPECS.get(name, {})
        weight_gb = cfg["model"].get("vram_estimate_gb", 0)
        display = spec.get("display", name)
        kv_budget = max(0.0, USABLE_VRAM_GB - weight_gb)
        if spec and spec.get("num_kv_layers"):
            try:
                kv_bpt = kv_bytes_per_token(spec, "BF16")
                kv_bpt_kb = kv_bpt / 1024
                s_bf16 = max_concurrent_sessions(weight_gb, kv_bpt, 4096)
                s_fp8 = max_concurrent_sessions(weight_gb, kv_bpt * 0.5, 4096)
                s_tq = (
                    max_concurrent_sessions(weight_gb, kv_bpt * 3 / 16, 4096)
                    if spec.get("turboquant_compatible")
                    else "N/A"
                )
                rows.append([
                    display, f"{weight_gb}", f"{kv_budget:.1f}",
                    f"{kv_bpt_kb:.1f}", str(s_bf16), str(s_fp8), str(s_tq),
                ])
            except Exception:
                rows.append([display, f"{weight_gb}", f"{kv_budget:.1f}", "—", "—", "—", "—"])
        else:
            rows.append([display, f"{weight_gb}", f"{kv_budget:.1f}", "—", "—", "—", "—"])
    return rows


# ---------------------------------------------------------------------------
# Entry point
# ---------------------------------------------------------------------------

def main() -> None:
    configs = []
    for yaml_path in sorted(CONFIGS_DIR.glob("*.yaml")):
        with yaml_path.open() as f:
            configs.append(yaml.safe_load(f))

    print(f"Loaded {len(configs)} Cat A model configs from {CONFIGS_DIR}")

    doc = build_report(configs)

    OUTPUT_PATH.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Report saved to {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
